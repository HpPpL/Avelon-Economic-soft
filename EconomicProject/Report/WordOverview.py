import os.path

import docx
import openpyxl
from docx.shared import Inches
from docx.shared import Pt
from ChartsToImage import ChartsToImage
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import time


class WordRecorder:
    def __init__(self, excel_file_path):
        # Разделим на две части - Excel и Word.
        self.time = time.time()
        # Excel:
        self.excel_file_path = excel_file_path
        self.excel_file_name = os.path.basename(excel_file_path)
        self.workbook = openpyxl.load_workbook(excel_file_path)
        self.sheets = self.workbook.sheetnames

        # Word:
        self.doc = docx.Document()
        self.filename = self.excel_file_name.split('.')[0] + ".docx"

    @staticmethod
    def wrap_row_data(worksheet, index, start_column):
        for row in worksheet.iter_rows(min_row=index, max_row=index, min_col=start_column, values_only=True):
            row_text = "\t".join(str(cell) for cell in row)  # Разделение ячеек табуляцией
            return list(map(lambda x: round(float(x), 2), row_text.split()))

    @staticmethod
    def wrap_column_data(worksheet, column, num=10):
        return [cell[0].value for cell in worksheet.iter_rows(min_col=column, max_col=column)][1:num]

    @staticmethod
    def safe_round(value, num):
        if value is None:
            return ""
        else:
            return round(float(value), num)

    def create_picture(self, path='Charts'):
        ChartsToImage(self.workbook).charts_to_image()

    def create_npv_page(self):
        worksheet = self.workbook['NPV']
        npv_text = worksheet['A29'].value
        # Для начала создадим таблицу
        data = {
            'NPV': ['млн.руб', self.safe_round(worksheet['D13'].value, 2)],
            'IRR': ['доли ед.', self.safe_round(worksheet['D14'].value, 2)],
            'Индекс доходности': ['', self.safe_round(worksheet['D17'].value, 2)],
            'ПРОСТОЙ СРОК ОКУПАЕМОСТИ (ГОД)': ['год', self.safe_round(worksheet['D18'].value, 2)],
            'ДИСКОНТИРОВАННЫЙ СРОК ОКУПАЕМОСТИ (ГОД)': ['год', self.safe_round(worksheet['D19'].value, 2)]
        }

        table = self.doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Добавление данных в таблицу
        for key, values in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = values[0]
            row_cells[2].text = str(values[1])

        # Объединим ячейку
        row = table.rows[2]
        cell1, cell2 = row.cells[:2]
        cell1.merge(cell2)

        # Создаем текст вывода
        self.doc.add_heading('Анализ критериев эффективности', 1)
        self.doc.add_paragraph(npv_text)

        # Добавляем картинку
        self.doc.add_picture('Charts/график1.png', width=docx.shared.Cm(16))

        # Добавляем разрыв страницы
        self.doc.add_page_break()

    def create_elastic_page(self):
        self.doc.add_heading('Анализ риска', 1)

        # Добавляем изображения
        self.doc.add_picture('Charts/график2.png', width=docx.shared.Cm(16))
        self.doc.add_picture('Charts/график3.png', width=docx.shared.Cm(16))
        self.doc.add_picture('Charts/график4.png', width=docx.shared.Cm(16))

    def create_montecarlo_page(self):
        worksheet = self.workbook['Credits']
        data = {
            'Метод': ['', worksheet['C15'].value],
            'ПОСТУПЛЕНИЕ': ['млн.руб', self.safe_round(worksheet['D16'].value, 2)],
            'Момент выдачи': ['год', self.safe_round(worksheet['D17'].value, 2)],
            'ЛЬГОТНЫЙ ПЕРИОД': ['год', self.safe_round(worksheet['D18'].value, 2)],
            'ПРОЦЕНТНАЯ СТАВКА': ['доли ед.', self.safe_round(worksheet['D19'].value, 2)],
            'ДЛИТЕЛЬНОСТЬ ЗАЙМА': ['год', self.safe_round(worksheet['D20'].value, 2)],
            'КАПИТАЛИЗАЦИЯ ПРОЦЕНТОВ': ['год', self.safe_round(worksheet['D21'].value, 2)]
        }

        self.doc.add_picture('Charts/график5.png', width=docx.shared.Cm(16))
        self.doc.add_heading('Анализ проектного финансирования', 1)

        table = self.doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Добавление данных в таблицу
        for key, values in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = values[0]
            row_cells[2].text = str(values[1])

        # Объединим ячейку
        row = table.rows[0]
        cell1, cell2 = row.cells[1:3]
        cell1.merge(cell2)

    def create_header(self):
        sections = self.doc.sections
        for section in sections:
            header = section.header
            if not header:
                header = section.add_header()

            paragraph = header.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
            run = paragraph.add_run()
            run.add_text('Итоговый сводный отчет по эффективности проекта')

    def make_style(self):
        # задаем стиль текста по умолчанию
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        self.doc.add_paragraph('Текст документа')

    def record(self):
        self.create_picture()
        self.create_npv_page()
        self.create_elastic_page()
        self.create_montecarlo_page()

        self.create_header()
        self.save_document()

    def save_document(self, path="Results/"):
        self.doc.save(path + self.filename)

    def __del__(self):
        print(f"Файл {self.filename} закрыт!")
        print(f"Время обработки составило {round(time.time() - self.time, 2)} с.")


if __name__ == "__main__":
    excel_path = "Results/Overview_E1.xlsx"
    dc = WordRecorder(excel_path)
    dc.record()
