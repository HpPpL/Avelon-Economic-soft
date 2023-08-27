import os.path

import docx
import openpyxl
from docx.shared import Inches
from docx.shared import Pt


class WordRecorder:
    def __init__(self, excel_file_path):
        # Разделим на две части - Excel и Word.
        # Excel:
        self.excel_file_path = excel_file_path
        self.excel_file_name = os.path.basename(excel_file_path)
        self.workbook = openpyxl.load_workbook(excel_file_path)
        self.sheets = self.workbook.sheetnames

        # Word:
        self.doc = docx.Document()
        self.filename = self.excel_file_name.split('.')[0] + ".docx"

    def make_style(self):
        # задаем стиль текста по умолчанию
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        self.doc.add_paragraph('Текст документа')

    def create_picture(self, path='Charts'):
        # for row in sheet.iter_rows(min_row=1, values_only=True):
        #     row_text = "\t".join(str(cell) for cell in row)  # Разделение ячеек табуляцией
        #     print(type(row_text))

            print()
            # doc.add_paragraph(row_text)

    def record(self):
        self.save_document()

    def save_document(self, path="Result/"):
        self.doc.save(path + self.filename)

    def __del__(self):
        print(f"Файл {self.filename} закрыт!")


if __name__ == "__main__":
    excel_path = "./Record_Report_E3.xlsx"
    dc = WordRecorder(excel_path)
    dc.save_document()
