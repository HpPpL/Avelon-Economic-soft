# import openpyxl
# import matplotlib.pyplot as plt
# from docx import Document
# from docx.shared import Inches
# from openpyxl_image_loader import SheetImageLoader
# # Откройте файл Excel
# excel_file_path = "Record.xlsx"
# wb = openpyxl.load_workbook(excel_file_path)
#
# # Создайте новый документ Word
# doc = Document()
# doc.add_heading('Графики из Excel', level=1)
#
# # Проход по листам Excel и поиск графиков
# for sheet in wb.sheetnames:
#     ws = wb[sheet]
#     for chart in ws._charts:
#         # Создание временного изображения графика с помощью matplotlib
#         temp_chart_image_path = 'temp_chart_image.png'
#         plt.plot(chart)  # Создаем пустой график для заполнения данными
#         # chart.refresh()
#         plt.savefig(temp_chart_image_path)
#         plt.close()
#
#         # Добавление изображения графика в документ Word
#         doc.add_paragraph(f"График на листе '{sheet}'")
#         doc.add_picture(temp_chart_image_path, width=Inches(6))
#
# # Сохранение документа Word
# output_doc_path = 'графики.docx'
# doc.save(output_doc_path)



from openpyxl import load_workbook
from docx import Document

# Загрузка книги Excel и выбор нужного листа
excel_file = "Record_Report_E3.xlsx"
wb = load_workbook(excel_file)
sheet = wb["Record"]

# Создание нового документа Word
doc = Document()

# Чтение данных из листа Excel и добавление их в документ Word
for row in sheet.iter_rows(min_row=1, values_only=True):
    row_text = "\t".join(str(cell) for cell in row)  # Разделение ячеек табуляцией
    print(row_text.split())

    print()
    # doc.add_paragraph(row_text)

# Сохранение документа Word
word_file = "результат.docx"
doc.save(word_file)
